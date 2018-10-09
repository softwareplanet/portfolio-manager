from django.core.exceptions import ObjectDoesNotExist
from rest_framework import status, permissions
from rest_framework.response import Response
from rest_framework.views import APIView

from employee.model_views import MultipleInstanceAPIView, SingleInstanceAPIView, \
    MultipleEmployeeRelatedInstanceAPIView, SingleEmployeeRelatedInstanceAPIView
from employee.models import Employee, Project, EmployeeProject, School, Skill, EmployeeSkill, EmployeeSchool, \
    SkillCategory, ProjectFile, FilesGroup
from employee.permissions import IsAdminOrSelf, IsPostOrIsAdmin
from employee.search_serializers import SearchEmployeeSerializer, SearchProjectSerializer, SearchSkillSerializer
from employee.serializers import EmployeeSerializer, EmployeeForUserSerializer, ProjectSerializer, \
    EmployeeProjectSerializer, SchoolSerializer, \
    SkillSerializer, EmployeeSkillSerializer, EmployeeSchoolSerializer, ExtendedProjectSerializer, \
    ExtendedSkillSerializer, ExtendedSchoolSerializer, SkillCategorySerializer, ChangePasswordSerializer, \
    ProjectFileSerializer, FilesGroupSerializer
from employee.utils import Utils


class ListEmployees(MultipleInstanceAPIView):
    serializer = EmployeeSerializer
    model = Employee
    permission_classes = (IsPostOrIsAdmin,)

    def get(self, request):
        models = self.model.objects.filter(is_active=1)
        return Response(self.serializer(models, many=True).data)


class ListEmployee(SingleInstanceAPIView):
    serializer = EmployeeSerializer
    serializer_for_user = EmployeeForUserSerializer
    model = Employee
    permission_classes = (permissions.IsAuthenticated, IsAdminOrSelf)

    def delete(self, request, model_id):
        try:
            model = self.model.objects.get(id=model_id)
            model.is_active = False
            model.save()
            return Response({'id': model_id}, status.HTTP_200_OK)
        except ObjectDoesNotExist as e:
            return Utils.error_response(e.args, status.HTTP_404_NOT_FOUND)


class ListProjects(MultipleInstanceAPIView):
    serializer = ProjectSerializer
    model = Project


class ListProjectsExtended(MultipleInstanceAPIView):
    serializer = ExtendedProjectSerializer
    model = Project


class ListFilesGroups(MultipleInstanceAPIView):
    serializer = FilesGroupSerializer
    model = FilesGroup


class ListProjectFiles(APIView):
    serializer = ProjectFileSerializer
    model = ProjectFile

    def get(self, request, model_id):
        models = self.model.objects.filter(project_id=model_id)
        return Response(self.serializer(models, many=True).data)

    def post(self, request, model_id):
        data = request.data
        data['project'] = model_id
        serializer = self.serializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Utils.error_response(serializer.errors, status.HTTP_400_BAD_REQUEST)


class ListProjectFile(SingleInstanceAPIView):
    serializer = ProjectFileSerializer
    model = ProjectFile
    permission_classes = (permissions.IsAdminUser,)


class ListProject(SingleInstanceAPIView):
    serializer = ExtendedProjectSerializer
    model = Project


class ListSkills(MultipleInstanceAPIView):
    serializer = SkillSerializer
    model = Skill

    def get(self, request):
        skills = self.model.objects.all()
        categories = SkillCategory.objects.all()
        return Response({
            'skills': self.serializer(skills, many=True).data,
            'categories': SkillCategorySerializer(categories, many=True).data
        })


class ListSkill(SingleInstanceAPIView):
    serializer = ExtendedSkillSerializer
    model = Skill


class ListSchools(MultipleInstanceAPIView):
    serializer = SchoolSerializer
    model = School


class ListSchool(SingleInstanceAPIView):
    serializer = ExtendedSchoolSerializer
    model = School


class ListEmployeeProjects(MultipleEmployeeRelatedInstanceAPIView):
    serializer = EmployeeProjectSerializer
    model = EmployeeProject

    def post(self, request, employee_id):
        if self._owner_or_admin(request, employee_id):
            data = request.data
            data['employeeId'] = employee_id
            serializer = self.serializer(data=data)
            if serializer.is_valid():
                serializer.save()
                self.__process_skills(serializer.data['skills'], employee_id)
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Utils.error_response(serializer.errors, status.HTTP_400_BAD_REQUEST)
        else:
            return Utils.error_response("Permission denied", status.HTTP_403_FORBIDDEN)

    @staticmethod
    def __process_skills(skills, user_id):
        for skill in skills:
            try:
                EmployeeSkill.objects.get(skill_id_id=skill['id'], employee_id_id=user_id)
            except ObjectDoesNotExist:
                EmployeeSkill.objects.create(employee_id_id=user_id, skill_id_id=skill['id'], level=1)


class ListEmployeeProject(SingleEmployeeRelatedInstanceAPIView):
    serializer = EmployeeProjectSerializer
    model = EmployeeProject

    def patch(self, request, employee_id, model_id):
        try:
            model = self.model.objects.get(id=model_id, employee_id=employee_id)
            if self._owner_or_admin(request, employee_id):
                serializer = self.serializer(model, data=request.data, partial=True)
                if serializer.is_valid():
                    serializer.save()
                    self.__process_skills(serializer.data['skills'], employee_id)
                    return Response(serializer.data, status.HTTP_200_OK)
                else:
                    return Utils.error_response(serializer.errors, status.HTTP_400_BAD_REQUEST)
            else:
                return Utils.error_response("Permission denied", status.HTTP_403_FORBIDDEN)
        except ObjectDoesNotExist as e:
            return Utils.error_response(e.args, status.HTTP_404_NOT_FOUND)

    @staticmethod
    def __process_skills(skills, user_id):
        for skill in skills:
            try:
                EmployeeSkill.objects.get(skill_id_id=skill['id'], employee_id_id=user_id)
            except ObjectDoesNotExist:
                EmployeeSkill.objects.create(employee_id_id=user_id, skill_id_id=skill['id'], level=1)


class ListEmployeeSkills(MultipleEmployeeRelatedInstanceAPIView):
    serializer = EmployeeSkillSerializer
    model = EmployeeSkill

    def post(self, request, employee_id):
        if self._owner_or_admin(request, employee_id):
            data = request.data
            data['employeeId'] = employee_id
            try:
                EmployeeSkill.objects.get(skill_id=data['skillId'], employee_id=data['employeeId'])
                return Utils.error_response({'non_field_errors': ['You already have such skill']},
                                            status.HTTP_400_BAD_REQUEST)
            except ObjectDoesNotExist:
                serializer = self.serializer(data=data)
                if serializer.is_valid():
                    serializer.save()
                    return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Utils.error_response(serializer.errors, status.HTTP_400_BAD_REQUEST)
        else:
            return Utils.error_response("Permission denied", status.HTTP_403_FORBIDDEN)


class ListEmployeeSkill(SingleEmployeeRelatedInstanceAPIView):
    serializer = EmployeeSkillSerializer
    model = EmployeeSkill

    def delete(self, request, employee_id, model_id):
        try:
            model = self.model.objects.get(id=model_id, employee_id=employee_id)
            if self._owner_or_admin(request, employee_id):
                for project in EmployeeProject.objects.filter(employee_id=employee_id):
                    project.skills.remove(model.skill_id)
                model.delete()
                return Response({'id': model_id}, status.HTTP_200_OK)
            else:
                return Utils.error_response("Permission denied", status.HTTP_403_FORBIDDEN)
        except ObjectDoesNotExist as e:
            return Utils.error_response(e.args, status.HTTP_404_NOT_FOUND)


class ListEmployeeSchools(MultipleEmployeeRelatedInstanceAPIView):
    serializer = EmployeeSchoolSerializer
    model = EmployeeSchool


class ListEmployeeSchool(SingleEmployeeRelatedInstanceAPIView):
    serializer = EmployeeSchoolSerializer
    model = EmployeeSchool


class ListMe(APIView):
    serializer = EmployeeForUserSerializer
    permission_classes = (permissions.IsAuthenticated,)

    def get(self, request):
        return Response(EmployeeForUserSerializer(request.user).data)

    def patch(self, request):
        try:
            model = request.user
            serializer = self.serializer(model, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status.HTTP_200_OK)
            else:
                return Utils.error_response(serializer.errors, status.HTTP_400_BAD_REQUEST)
        except ObjectDoesNotExist as e:
            return Utils.error_response(e.args, status.HTTP_404_NOT_FOUND)

    def put(self, request):
        user = request.user
        serializer = ChangePasswordSerializer(data=request.data)

        if serializer.is_valid():
            old_password = serializer.data.get("old_password")
            if not user.check_password(old_password):
                return Utils.error_response({"old_password": ["Wrong password."]}, status.HTTP_400_BAD_REQUEST)
            user.set_password(serializer.data.get("new_password"))
            user.save()
            return Response(status=status.HTTP_204_NO_CONTENT)

        return Utils.error_response(serializer.errors, status.HTTP_400_BAD_REQUEST)


class ListSearch(APIView):
    permission_classes = (permissions.IsAuthenticated,)

    def post(self, request):
        from django.db.models import Q
        result = []
        data = request.data['query']
        if request.user.is_staff:
            result.extend(SearchEmployeeSerializer(
                Employee.objects.filter(
                    Q(first_name__istartswith=data) | Q(last_name__istartswith=data) | Q(
                        description__icontains=(' ' + data)) | Q(description__istartswith=data)),
                many=True).data)
            result.extend(SearchProjectSerializer(
                Project.objects.filter(
                    Q(name__istartswith=data) | Q(description__icontains=(' ' + data))),
                many=True).data)
            result.extend(SearchSkillSerializer(
                Skill.objects.filter(
                    Q(name__istartswith=data)),
                many=True).data)
        else:
            pass
        return Response(result)


class ListDocCandidatePresentation(APIView):
    from django.http import HttpResponse
    from io import BytesIO
    employee = None
    candidate_presentation = None
    permission_classes = (permissions.AllowAny,)

    def get(self, request, model_id):
        from docx import Document
        try:
            self.employee = Employee.objects.get(id=model_id)
            self.candidate_presentation = Document()
            self.candidate_presentation.add_heading('Candidate Presentation', 0)
            self._generate_employee_summary()
            if self.employee.position is not None and self.employee.career_start_date is not None:
                self._generate_employee_summary_of_qualification()
            return self._generate_file_response()
        except ObjectDoesNotExist as e:
            return Utils.error_response(e.args, status.HTTP_404_NOT_FOUND)

    def _generate_employee_summary(self):
        photo = self.candidate_presentation.add_picture(self.employee.image.file)
        name = self.candidate_presentation.add_paragraph('Name: ')
        run = name.add_run(self.employee.get_full_name())
        run.bold = True
        position = self.candidate_presentation.add_paragraph('Position: {0}'.format(self.employee.position))
        description = self.candidate_presentation.add_paragraph(self.employee.description)

    def _generate_employee_summary_of_qualification(self):
        from datetime import date
        import math
        self.candidate_presentation.add_heading('Summary of qualification')
        experience = self.candidate_presentation.add_paragraph('{0} years of experience as {1}'.format(
            math.ceil((date.today() - self.employee.career_start_date).days / 365.0),
            self.employee.position))

    def _generate_file_response(self):
        f = self.BytesIO()
        self.candidate_presentation.save(f)
        length = f.tell()
        f.seek(0)
        response = self.HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename={0}.docx'.format(self.employee.get_full_name())
        response['Content-Length'] = length
        return response
